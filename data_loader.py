from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document


ARABIC_DIACRITICS = re.compile(r"[\u0610-\u061a\u064b-\u065f\u0670\u06d6-\u06ed]")
ARABIC_DIRECT_DIACRITICS = re.compile(r"[\u0610-\u061a\u064b-\u065f\u0670\u06d6-\u06ed]")
STORY_FILE_PATTERN = re.compile(r"^(?P<number>\d+)\s*-\s*(?P<title>.+?)\.docx$", re.IGNORECASE)
IGNORED_DOCX_PREFIXES = ("~$", "000-", "Quality-Control")

DETAIL_ALIASES = {
    "هود": ["عاد"],
    "صالح": ["ثمود"],
    "اسماعيل": ["الذبيح"],
    "الكهف": ["اصحاب الكهف"],
}

SECTION_LABELS = {
    "overview": {"المحور المركزي", "السؤال الإنساني", "الحركة الكبرى"},
    "applications": {"المجال", "التطبيق"},
    "stages": {"المرحلة", "النص أو المعنى المحوري", "النص", "التطبيق على الإنسان"},
    "review": {"المراجعة", "المعيار", "الإنجاز", "لماذا لم يصل إلى 100%؟", "الإجراء المتبقي"},
    "cycle": {"الدورة", "ما تم فحصه وتحسينه", "الأثر"},
}


@dataclass(frozen=True)
class StoryIndexItem:
    id: str
    number: int
    title: str
    character: str
    subtitle: str
    category: str
    core_value: str
    surah_references: str
    significance: str
    skills: str
    docx_path: str | None
    source: str


@dataclass(frozen=True)
class StoryContent:
    id: str
    number: int
    title: str
    subtitle: str
    category: str
    core_value: str
    surah_references: str
    significance: str
    skills: list[str]
    verses: list[dict[str, Any]]
    overview: list[dict[str, str]]
    summary: str
    transformation_chain: list[dict[str, str]]
    practical_applications: dict[str, str]
    reflections: list[str]
    turning_points: list[str]
    source_file: str | None


def normalize_arabic(value: Any) -> str:
    text = "" if value is None else str(value)
    text = ARABIC_DIACRITICS.sub("", text.strip())
    text = text.replace("ـ", "")
    text = re.sub("[إأآٱا]", "ا", text)
    text = text.replace("ى", "ي").replace("ة", "ه").replace("ؤ", "و").replace("ئ", "ي")
    text = re.sub(r"[^\w\s\u0600-\u06ff]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip().lower()


def normalize_search_text(value: Any) -> str:
    text = "" if value is None else str(value)
    text = ARABIC_DIRECT_DIACRITICS.sub("", text.strip())
    text = text.replace("ـ", "")
    text = re.sub("[إأآٱا]", "ا", text)
    text = text.replace("ى", "ي").replace("ة", "ه").replace("ؤ", "و").replace("ئ", "ي")
    text = re.sub(r"[^\w\s\u0600-\u06ff]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip().lower()


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def story_id(number: int) -> str:
    return f"story-{number:03d}"


def find_workbook(base_dir: Path) -> Path:
    candidates = [
        base_dir / "Stories included in holy (Quran).xlsx",
        base_dir.parent / ".streamlit" / "Stories included in holy (Quran).xlsx",
        base_dir.parent / "Stories included in holy (Quran).xlsx",
    ]
    candidates.extend(path for path in base_dir.glob("*.xlsx") if not path.name.startswith("~$"))
    streamlit_dir = base_dir.parent / ".streamlit"
    if streamlit_dir.exists():
        candidates.extend(path for path in streamlit_dir.glob("*.xlsx") if not path.name.startswith("~$"))
    for path in candidates:
        if path.exists() and not path.name.startswith("~$"):
            return path
    raise FileNotFoundError("Missing Excel workbook for Quran stories.")


def discover_story_docx(base_dir: Path) -> list[Path]:
    paths: list[Path] = []
    for path in base_dir.glob("*.docx"):
        if path.name.startswith(IGNORED_DOCX_PREFIXES):
            continue
        if STORY_FILE_PATTERN.match(path.name):
            paths.append(path)
    return sorted(paths, key=lambda p: int(STORY_FILE_PATTERN.match(p.name).group("number")))


def _docx_name_parts(path: Path) -> tuple[int, str]:
    match = STORY_FILE_PATTERN.match(path.name)
    if not match:
        raise ValueError(f"Unexpected story filename: {path.name}")
    return int(match.group("number")), match.group("title").strip()


def _candidate_keys(title: str) -> set[str]:
    keys = {normalize_arabic(title)}
    for key in list(keys):
        keys.update(normalize_arabic(alias) for alias in DETAIL_ALIASES.get(key, []))
    return {key for key in keys if key}


def _docx_key_map(base_dir: Path) -> dict[str, Path]:
    mapping: dict[str, Path] = {}
    for path in discover_story_docx(base_dir):
        _, title = _docx_name_parts(path)
        for key in _candidate_keys(title):
            mapping.setdefault(key, path)
    return mapping


def _read_docx_metadata(path: Path) -> dict[str, str]:
    number, filename_title = _docx_name_parts(path)
    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    title = filename_title
    subtitle = ""
    core_value = ""
    category = ""
    references = ""
    if len(paragraphs) >= 2 and paragraphs[0].isdigit():
        title = paragraphs[1]
        subtitle = paragraphs[2] if len(paragraphs) > 2 else ""
    elif paragraphs:
        title = paragraphs[0]
        subtitle = paragraphs[1] if len(paragraphs) > 1 else ""
    for paragraph in paragraphs[:8]:
        if paragraph.startswith("القيمة المحورية:"):
            parts = [part.strip() for part in paragraph.split("|")]
            core_value = parts[0].replace("القيمة المحورية:", "").strip()
            if len(parts) > 1:
                category = parts[1].replace("التصنيف:", "").strip()
        if paragraph.startswith("مواضع القصة:"):
            references = paragraph.replace("مواضع القصة:", "").strip()
    return {
        "number": str(number),
        "title": clean_text(title),
        "subtitle": clean_text(subtitle),
        "core_value": clean_text(core_value),
        "category": clean_text(category),
        "surah_references": clean_text(references),
    }


def load_story_index(base_dir: Path) -> list[dict[str, Any]]:
    workbook = find_workbook(base_dir)
    df = pd.read_excel(workbook, sheet_name="Overall Stories")
    df = df.loc[df["#"].notna()].copy()
    docx_keys = _docx_key_map(base_dir)
    items: list[StoryIndexItem] = []
    used_docx: set[Path] = set()

    for _, row in df.iterrows():
        number = int(row["#"])
        title = clean_text(row.get("Story", ""))
        character = clean_text(row.get("Character", ""))
        story_key = normalize_arabic(title)
        character_key = normalize_arabic(character)
        docx_path = docx_keys.get(story_key) or docx_keys.get(character_key)
        if docx_path:
            used_docx.add(docx_path)
        items.append(
            StoryIndexItem(
                id=story_id(number),
                number=number,
                title=title,
                character=character,
                subtitle=clean_text(row.get("1-line Summary", "")),
                category=clean_text(row.get("Category", "")),
                core_value=clean_text(row.get("Core Value", "")),
                surah_references=clean_text(row.get("السور والآيات", "")),
                significance=clean_text(row.get("Significance", "")),
                skills=clean_text(row.get("Skill to be gained", "")),
                docx_path=str(docx_path) if docx_path else None,
                source="excel",
            )
        )

    existing_ids = {item.id for item in items}
    for path in discover_story_docx(base_dir):
        if path in used_docx:
            continue
        metadata = _read_docx_metadata(path)
        number = int(metadata["number"])
        candidate_id = story_id(number)
        if candidate_id in existing_ids:
            candidate_id = f"docx-{number:03d}"
        items.append(
            StoryIndexItem(
                id=candidate_id,
                number=number,
                title=metadata["title"],
                character=metadata["title"],
                subtitle=metadata["subtitle"],
                category=metadata["category"],
                core_value=metadata["core_value"],
                surah_references=metadata["surah_references"],
                significance="",
                skills="",
                docx_path=str(path),
                source="docx",
            )
        )
        existing_ids.add(candidate_id)

    return [item.__dict__ for item in sorted(items, key=lambda item: (item.number, item.id))]


def _cell_text(cell: Any) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip()).strip()


def _is_method_text(text: str) -> bool:
    return text.startswith("منهج العمل:")


def _is_verse_text(text: str) -> bool:
    head = text[:220]
    return (
        "سورة" in head
        or head.startswith("- سورة")
        or "﴿" in text
        or "الآيات من" in head
        or "الآية" in head
    )


def _looks_like_summary(text: str) -> bool:
    if len(text) < 160:
        return False
    verse_markers = ["﴿", "الآيات من", "سورة", "وَ", "قالَ", "قال "]
    return not any(marker in text[:120] for marker in verse_markers)


def _split_verse_and_explanation(text: str) -> tuple[str, list[str]]:
    if not _is_verse_text(text):
        return text, []
    parts = [part.strip() for part in re.split(r"\s*\|\s*|\n+", text) if part.strip()]
    if len(parts) < 3:
        return text, []
    verse_parts: list[str] = []
    explanation_parts: list[str] = []
    explanation_started = False
    for part in parts:
        if re.match(r"^\d+[\.\)]?\s+", part) and not _is_verse_text(part):
            explanation_started = True
        if explanation_started:
            explanation_parts.append(part)
        else:
            verse_parts.append(part)
    if not explanation_parts and "|" in text:
        verse_parts = parts[:-1]
        explanation_parts = parts[-1:]
    if not explanation_parts:
        return text, []
    return " | ".join(verse_parts).strip(), _split_explanation_text("\n".join(explanation_parts))


def _split_explanation_text(text: str) -> list[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if len(lines) > 1:
        return lines
    matches = re.findall(r"(?:^|\s)(\d+\s+.*?)(?=\s+\d+\s+|$)", text, flags=re.DOTALL)
    if matches:
        return [re.sub(r"\s+", " ", item).strip() for item in matches]
    cleaned = clean_text(text)
    return [cleaned] if cleaned else []


def _parse_story_docx(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    metadata = _read_docx_metadata(path)
    overview: list[dict[str, str]] = []
    verses: list[dict[str, Any]] = []
    transformation_chain: list[dict[str, str]] = []
    practical_applications: dict[str, str] = {}
    summary_candidates: list[str] = []
    reflections: list[str] = []

    for paragraph in paragraphs:
        if paragraph.isdigit() or paragraph in {"الآيات + الشرح:", "الايات + الشرح:", "الخلاصة:"}:
            continue
        if _is_method_text(paragraph) or paragraph.startswith("القيمة المحورية:") or paragraph.startswith("مواضع القصة:"):
            continue
        parts = [part.strip() for part in paragraph.split("|")]
        if len(parts) >= 3 and parts[0] != "المرحلة":
            transformation_chain.append(
                {"stage": parts[0], "text": parts[1], "application": " | ".join(parts[2:])}
            )
            continue
        if _looks_like_summary(paragraph):
            summary_candidates.append(paragraph)
        elif len(paragraph) > 30:
            reflections.append(paragraph)

    for table in document.tables:
        rows = [[_cell_text(cell) for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        first_row = rows[0]
        first_cell = first_row[0].strip() if first_row else ""

        if first_cell == "المحور المركزي" and len(first_row) >= 2:
            for row in rows:
                if len(row) >= 2 and row[0].strip() and row[1].strip():
                    overview.append({"label": row[0].strip(), "value": row[1].strip()})
            continue

        if first_cell == "المرحلة":
            for row in rows[1:]:
                if len(row) >= 3 and any(cell.strip() for cell in row):
                    transformation_chain.append(
                        {"stage": row[0].strip(), "text": row[1].strip(), "application": row[2].strip()}
                    )
            continue

        if first_cell == "المجال":
            for row in rows[1:]:
                if len(row) >= 2 and row[0].strip() and row[1].strip():
                    practical_applications[row[0].strip()] = row[1].strip()
            continue

        if first_cell in {"المراجعة", "الدورة"}:
            continue

        if len(first_row) == 1:
            flat_rows = [row[0].strip() for row in rows if row and row[0].strip()]
            if flat_rows and len(flat_rows[0]) <= 80 and not _is_verse_text(flat_rows[0]):
                flat_rows = flat_rows[1:]
            pending_verse = ""
            for text in flat_rows:
                if _is_method_text(text):
                    continue
                verse, explanation = _split_verse_and_explanation(text)
                if explanation:
                    verses.append({"text": verse, "explanation": explanation})
                    pending_verse = ""
                    continue
                if _is_verse_text(text):
                    if pending_verse:
                        verses.append({"text": pending_verse, "explanation": []})
                    pending_verse = text
                    continue
                if pending_verse:
                    verses.append({"text": pending_verse, "explanation": _split_explanation_text(text)})
                    pending_verse = ""
                    continue
                if _looks_like_summary(text):
                    summary_candidates.append(text)
            if pending_verse:
                verses.append({"text": pending_verse, "explanation": []})

    summary = summary_candidates[0] if summary_candidates else ""
    return {
        "metadata": metadata,
        "verses": verses,
        "overview": overview,
        "summary": summary,
        "transformation_chain": transformation_chain,
        "practical_applications": practical_applications,
        "reflections": reflections,
        "turning_points": [row["stage"] for row in transformation_chain if row.get("stage")],
    }


def _split_skills(skills: str) -> list[str]:
    return [part.strip() for part in re.split(r"[-–—،,]+", skills) if part.strip()]


class StoryRepository:
    def __init__(self, base_dir: str | Path):
        self.base_dir = Path(base_dir)

    def get_index(self) -> list[dict[str, Any]]:
        return load_story_index(self.base_dir)

    def get_story(self, story_id_value: str) -> dict[str, Any] | None:
        index = self.get_index()
        item = next((story for story in index if story["id"] == story_id_value), None)
        if not item:
            return None
        parsed = _parse_story_docx(Path(item["docx_path"])) if item.get("docx_path") else {}
        metadata = parsed.get("metadata", {})
        title = item.get("title") or metadata.get("title", "")
        subtitle = item.get("subtitle") or metadata.get("subtitle", "")
        category = item.get("category") or metadata.get("category", "")
        core_value = item.get("core_value") or metadata.get("core_value", "")
        surah_references = item.get("surah_references") or metadata.get("surah_references", "")
        content = StoryContent(
            id=item["id"],
            number=int(item["number"]),
            title=title,
            subtitle=subtitle,
            category=category,
            core_value=core_value,
            surah_references=surah_references,
            significance=item.get("significance", ""),
            skills=_split_skills(item.get("skills", "")),
            verses=parsed.get("verses", []),
            overview=parsed.get("overview", []),
            summary=parsed.get("summary", ""),
            transformation_chain=parsed.get("transformation_chain", []),
            practical_applications=parsed.get("practical_applications", {}),
            reflections=parsed.get("reflections", []),
            turning_points=parsed.get("turning_points", []),
            source_file=Path(item["docx_path"]).name if item.get("docx_path") else None,
        )
        return content.__dict__

    def get_previous_next(self, story_id_value: str) -> tuple[str | None, str | None]:
        index = self.get_index()
        ids = [story["id"] for story in index]
        if story_id_value not in ids:
            return None, None
        position = ids.index(story_id_value)
        previous_id = ids[position - 1] if position > 0 else None
        next_id = ids[position + 1] if position < len(ids) - 1 else None
        return previous_id, next_id

    def get_related(self, story_id_value: str) -> list[dict[str, Any]]:
        index = self.get_index()
        current = next((story for story in index if story["id"] == story_id_value), None)
        if not current:
            return []
        related = [
            story
            for story in index
            if story["id"] != story_id_value
            and (
                story.get("category") == current.get("category")
                or story.get("core_value") == current.get("core_value")
            )
        ]
        return related[:4]

    def search(
        self,
        query: str = "",
        category: str | None = None,
        value: str | None = None,
        skill: str | None = None,
    ) -> list[dict[str, Any]]:
        query_value = normalize_search_text(query)
        category_value = "" if category in {None, "", "all", "الكل"} else normalize_search_text(category)
        value_value = "" if value in {None, "", "all", "الكل"} else normalize_search_text(value)
        skill_value = "" if skill in {None, "", "all", "الكل"} else normalize_search_text(skill)
        results: list[dict[str, Any]] = []
        for story in self.get_index():
            searchable = normalize_search_text(
                " ".join(
                    [
                        story.get("title", ""),
                        story.get("character", ""),
                        story.get("subtitle", ""),
                        story.get("category", ""),
                        story.get("core_value", ""),
                        story.get("surah_references", ""),
                        story.get("significance", ""),
                        story.get("skills", ""),
                    ]
                )
            )
            if query_value and query_value not in searchable:
                continue
            if category_value and category_value != normalize_search_text(story.get("category", "")):
                continue
            if value_value and value_value != normalize_search_text(story.get("core_value", "")):
                continue
            if skill_value and skill_value not in normalize_search_text(story.get("skills", "")):
                continue
            results.append(story)
        return results
